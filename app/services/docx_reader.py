# app/services/docx_reader.py
from typing import Dict, Any, List, Tuple
from io import BytesIO
from docx import Document
import re

# ---------- helpers ----------
def _t(s: str | None) -> str:
    return (s or "").strip()

def _kv(line: str) -> Tuple[str, str] | None:
    """
    Convierte líneas tipo "Clave: Valor" en (clave, valor).
    Mantiene mayúsculas/minúsculas de la clave tal y como aparecen.
    """
    m = re.match(r"^\s*([^:]+)\s*:\s*(.*)$", line)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return None

def _collect_until(next_idx: int, paras: List[str], stop_markers: List[str]) -> Tuple[str, int]:
    """
    Acumula texto desde next_idx hasta que encuentre un párrafo que empiece
    por alguno de los stop_markers (prefijo case-insensitive) o se acaben los párrafos.
    Devuelve (texto_unido, nuevo_idx)
    """
    buf: List[str] = []
    i = next_idx
    while i < len(paras):
        p = paras[i].strip()
        low = p.lower()
        if any(low.startswith(m) for m in stop_markers):
            break
        buf.append(p)
        i += 1
    return "\n".join(buf).strip(), i

def _split_list(raw: str) -> List[str]:
    if not raw:
        return []
    # separa por coma o ; y quita espacios
    parts = [x.strip() for x in re.split(r"[;,]", raw) if x.strip()]
    # de-dup preservando orden
    seen = set(); out=[]
    for x in parts:
        if x not in seen:
            seen.add(x); out.append(x)
    return out

# ---------- parser principal ----------
def extract_fields_from_docx(docx_bytes: bytes) -> Dict[str, Any]:
    """
    Lee la ficha DOCX (plantilla SI) y devuelve un dict con claves que usa el transformer:
    - "Nombre de la ayuda"
    - "Portales"                  -> lista de strings
    - "Tipo de ayuda"             -> lista (máx 3 luego en transformer)
    - "Fecha inicio"
    - "Fecha fin"
    - "Fecha publicación en la BDNS o en el Boletín Oficial"
    - "Ámbito territorial"
    - "Administración"
    - "Plazo de presentación"
    - Bloques largos (texto libre):
      "Beneficiarios/Destinatarios", "Requisitos de acceso", "Descripción",
      "Cuantía", "Importe máximo", "Resolución", "Documentos a presentar",
      "Normativa Reguladora", "Referencia Legislativa", "Lugar y forma de presentación",
      "Costes no Subvencionables"
    - Otros datos (pie de ficha):
      "Usuario"   (USUARIO: ...)
      "Fecha"     (FECHA: ...)
      "Frase para publicitar" (opcional)
    """
    doc = Document(BytesIO(docx_bytes))

    # Convertimos todos los párrafos a texto lineal limpio
    paras = [p.text for p in doc.paragraphs]
    # Algunas plantillas ponen secciones en tablas; recogemos también celdas
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                # Evita duplicar párrafos si ya están
                t = cell.text.strip()
                if t and t not in paras:
                    paras.append(t)

    # Normalizamos espacios
    paras = [re.sub(r"\s+", " ", p).strip() for p in paras if p and p.strip()]
    # A veces la ficha usa bloque con títulos -> nos apoyamos en prefijos
    # Mapa de campos "línea simple"
    simple_map = {
        "Nombre de la ayuda": None,
        "Portales": None,
        "Categoría": None,
        "Tipo de ayuda": None,
        "Fecha inicio": None,
        "Fecha fin": None,
        "Fecha publicación en la BDNS o en el Boletín Oficial": None,
        "Ámbito territorial": None,
        "Administración": None,
        "Plazo de presentación": None,
    }

    # Bloques largos (pueden ocupar varias líneas)
    block_titles = [
        "Beneficiarios/Destinatarios",
        "Requisitos de acceso",
        "Descripción",
        "Cuantía",
        "Importe máximo",
        "Resolución",
        "Documentos a presentar",
        "Normativa Reguladora",
        "Referencia Legislativa",
        "Lugar y forma de presentación",
        "Costes no Subvencionables",
    ]
    block_results: Dict[str, str] = {k: "" for k in block_titles}

    # Otros datos al final
    otros_usuario = ""
    otros_fecha = ""
    otros_frase = ""

    # Índice para recorrido
    i = 0
    L = len(paras)
    # Prepara marcadores de parada para bloques
    block_stops = [x.lower() for x in (list(simple_map.keys()) + block_titles + ["Otros datos", "OTROS DATOS"])]

    while i < L:
        line = paras[i]
        low = line.lower()

        # 1) Otros datos
        if low.startswith("otros datos"):
            i += 1
            # Leemos hasta el final o hasta que aparezca un título conocido (poco probable)
            text, i = _collect_until(i, paras, block_stops)
            # Extrae USUARIO:, FECHA:, FRASE PARA PUBLICITAR:
            m_user = re.search(r"usuario\s*:\s*(.+)", text, flags=re.I)
            if m_user: otros_usuario = m_user.group(1).strip()
            m_fecha = re.search(r"fecha\s*:\s*([0-9]{1,2}/[0-9]{1,2}/[0-9]{2,4})", text, flags=re.I)
            if m_fecha: otros_fecha = m_fecha.group(1).strip()
            m_frase = re.search(r"frase\s+para\s+publicitar\s*:\s*(.+)", text, flags=re.I)
            if m_frase: otros_frase = m_frase.group(1).strip()
            continue

        # 2) Bloques largos (prefijo exacto)
        matched_block = next((t for t in block_titles if low.startswith(t.lower())), None)
        if matched_block:
            # Consumimos "Titulo:" si viene en la misma línea
            # y acumulamos párrafos hasta el siguiente título
            # Si la línea ya tiene "Titulo: contenido" lo contamos también
            content = ""
            kv = _kv(line)
            if kv:
                # "Titulo: contenido..."
                content = kv[1]
                i += 1
            else:
                i += 1
            more, i = _collect_until(i, paras, block_stops)
            joined = "\n".join(x for x in [content, more] if x).strip()
            block_results[matched_block] = joined
            continue

        # 3) Simples "Clave: Valor"
        kvp = _kv(line)
        if kvp:
            k, v = kvp
            if k in simple_map:
                simple_map[k] = v.strip()
                i += 1
                continue

        i += 1

    # Ensamblar resultado
    result: Dict[str, Any] = {}
    result.update({k: _t(v) for k, v in simple_map.items() if v})

    # Campos con formatos especiales
    # Portales / Tipo de ayuda -> listas
    if "Portales" in result:
        result["Portales"] = _split_list(result["Portales"])
    if "Tipo de ayuda" in result:
        result["Tipo de ayuda"] = _split_list(result["Tipo de ayuda"])

    # Bloques largos
    result.update({k: v for k, v in block_results.items() if v})

    # Otros datos
    if otros_usuario:
        result["Usuario"] = otros_usuario
    if otros_fecha:
        result["Fecha"] = otros_fecha
    if otros_frase:
        result["Frase para publicitar"] = otros_frase

    return result
